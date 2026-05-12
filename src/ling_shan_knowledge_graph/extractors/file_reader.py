# extractors/file_reader.py
import os
import re
from typing import List, Tuple, Dict


class FileReader:
    """读取Word/Excel文件内容"""

    @staticmethod
    def read_docx(file_path: str) -> Tuple[str, List[Dict]]:
        """读取Word文档，返回文本和表格数据"""
        try:
            from docx import Document
            doc = Document(file_path)

            # 读取段落文本
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

            # 读取表格数据（重要！您的结构化数据在这里）
            tables_data = []
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_content.append(row_data)
                tables_data.append(table_content)

            return text, tables_data

        except ImportError:
            print("❌ 请安装 python-docx: pip install python-docx")
            return "", []

    @staticmethod
    def read_file(file_path: str) -> Tuple[str, List[Dict]]:
        """自动识别文件类型并读取"""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.docx':
            return FileReader.read_docx(file_path)
        else:
            return FileReader.read_txt(file_path), []

    @staticmethod
    def read_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()