from __future__ import annotations

import csv
import io
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Tuple

from docx import Document
from fastapi import UploadFile

from services.poi_repository import PoiRepository, normalize_point_type


FIELD_ALIASES = {
    "poi_id": {"poi_id", "id", "编号", "景点编号", "点位编号", "景点ID", "景点id", "POI编号", "POIID"},
    "name": {"name", "poi_name", "景点名称", "点位名称", "名称", "景点", "点位", "标题"},
    "aliases": {"aliases", "alias", "别名", "关键词", "搜索词"},
    "lat": {"lat", "latitude", "纬度"},
    "lng": {"lng", "lon", "longitude", "经度"},
    "point_type": {"point_type", "type", "类型", "点位类型", "景点类型"},
    "stay_minutes": {"stay_minutes", "duration", "建议停留", "建议停留分钟", "停留时间"},
    "intro": {"intro", "description", "desc", "简介", "介绍", "描述", "详细介绍"},
    "address": {"address", "地址", "位置", "具体位置"},
    "tags": {"tags", "标签", "适合人群", "特色", "游玩亮点"},
    "scenic_name": {"景区名称", "景区", "所属景区"},
    "params": {"建筑/景观参数", "建筑参数", "景观参数", "参数"},
    "function": {"核心功能", "功能"},
    "culture": {"文化内涵", "文化"},
    "opening_info": {"演艺/开放信息", "开放信息", "演艺信息", "开放时间"},
    "remark": {"备注", "说明"},
}


class PoiImportService:
    SUPPORTED_SUFFIXES = {".txt", ".md", ".json", ".csv", ".docx"}

    def __init__(self, repository: PoiRepository, tencent_map_client: Any | None = None) -> None:
        self.repository = repository
        self.tencent_map_client = tencent_map_client

    async def import_poi_file(
        self,
        file: UploadFile,
        scenic_id: str,
        scenic_name: str | None = None,
        overwrite: bool = True,
        use_tencent_geocode: bool = True,
        city: str | None = None,
        address_hint: str | None = None,
    ) -> Dict[str, Any]:
        filename = file.filename or "poi.txt"
        suffix = Path(filename).suffix.lower()

        if suffix not in self.SUPPORTED_SUFFIXES:
            supported = ", ".join(sorted(self.SUPPORTED_SUFFIXES))
            raise ValueError(f"不支持的文件类型: {suffix or 'unknown'}。当前仅支持 {supported}")

        raw_bytes = await file.read()
        if not raw_bytes:
            raise ValueError("上传文件内容为空")

        parsed = self.parse_uploaded_file(
            scenic_id=scenic_id,
            scenic_name=scenic_name,
            filename=filename,
            suffix=suffix,
            raw_bytes=raw_bytes,
            use_tencent_geocode=use_tencent_geocode,
            city=city,
            address_hint=address_hint,
        )

        return self.save_pois(
            scenic_id=scenic_id,
            scenic_name=parsed["scenic_name"],
            pois=parsed["pois"],
            overwrite=overwrite,
            invalid_items=parsed["invalid_items"],
            source_file=filename,
        )

    def parse_uploaded_file(
        self,
        scenic_id: str,
        scenic_name: str | None,
        filename: str,
        suffix: str,
        raw_bytes: bytes,
        use_tencent_geocode: bool = True,
        city: str | None = None,
        address_hint: str | None = None,
    ) -> Dict[str, Any]:
        text = raw_bytes.decode("utf-8-sig", errors="ignore") if suffix != ".docx" else ""

        if suffix == ".docx":
            raw_pois, invalid_items, extracted_name = self.parse_docx(raw_bytes)
            scenic_name = scenic_name or extracted_name
        elif suffix == ".json":
            raw_pois, invalid_items, extracted_name = self.parse_json(text)
            scenic_name = scenic_name or extracted_name
        elif suffix == ".csv":
            raw_pois, invalid_items = self.parse_csv(text)
        elif suffix in {".txt", ".md"}:
            raw_pois, invalid_items = self.parse_txt_or_md(text, source="text_rule")
        else:
            raise ValueError(f"不支持的文件类型: {suffix}")

        valid_pois: List[Dict[str, Any]] = []
        collected_invalid = list(invalid_items)

        for index, raw_poi in enumerate(raw_pois, start=1):
            normalized = self.normalize_poi(
                raw=raw_poi,
                scenic_id=scenic_id,
                scenic_name=scenic_name,
                index=index,
                use_tencent_geocode=use_tencent_geocode,
                city=city,
                address_hint=address_hint,
            )
            if normalized.get("valid"):
                valid_pois.append(normalized["poi"])
            else:
                collected_invalid.append(normalized["invalid"])

        if not valid_pois:
            sample_reasons = collected_invalid[:10]
            raise ValueError(
                f"文件 {filename} 未解析出任何有效 POI，请检查表格字段、名称和坐标/地址信息。"
                f"解析失败样例: {sample_reasons}"
            )

        return {
            "scenic_id": scenic_id,
            "scenic_name": scenic_name,
            "pois": valid_pois,
            "invalid_items": collected_invalid,
        }

    # ============================================================
    # DOCX 解析增强版
    # 支持：
    # 1. 标准横向表格：第一行表头，后续每行一个 POI
    # 2. 纵向表格：第一列是字段名，后面每列一个 POI
    # 3. 你当前这种 docx：表头行 + 数据按“字段数量”循环展开
    # ============================================================

    def parse_docx(self, raw_bytes: bytes) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], str | None]:
        document = Document(io.BytesIO(raw_bytes))

        table_pois, table_invalid = self.parse_docx_tables(document)
        if table_pois:
            scenic_name = self._guess_scenic_name_from_docx(document)
            return table_pois, table_invalid, scenic_name

        paragraph_pois, paragraph_invalid = self.parse_docx_paragraphs(document)
        if paragraph_pois:
            scenic_name = self._guess_scenic_name_from_docx(document)
            return paragraph_pois, table_invalid + paragraph_invalid, scenic_name

        reasons = table_invalid + paragraph_invalid
        detail = reasons[0]["reason"] if reasons else "docx 中没有识别到表格或可用段落"
        raise ValueError(f"docx 没有解析出任何有效 POI: {detail}")

    def parse_docx_tables(self, document: Document) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        raw_pois: List[Dict[str, Any]] = []
        invalid_items: List[Dict[str, Any]] = []

        for table_index, table in enumerate(document.tables):
            matrix = self._table_to_matrix(table)
            if not matrix:
                continue

            parsed_by_horizontal, invalid_horizontal = self._parse_horizontal_table(matrix, table_index)
            if parsed_by_horizontal:
                raw_pois.extend(parsed_by_horizontal)
                invalid_items.extend(invalid_horizontal)
                continue

            parsed_by_vertical, invalid_vertical = self._parse_vertical_table(matrix, table_index)
            if parsed_by_vertical:
                raw_pois.extend(parsed_by_vertical)
                invalid_items.extend(invalid_vertical)
                continue

            parsed_by_flat_blocks, invalid_flat_blocks = self._parse_flat_block_table(matrix, table_index)
            if parsed_by_flat_blocks:
                raw_pois.extend(parsed_by_flat_blocks)
                invalid_items.extend(invalid_flat_blocks)
                continue

            invalid_items.append(
                {
                    "table_index": table_index,
                    "reason": "表格不是标准横向表格、纵向表格，也未识别出按字段循环排列的数据块",
                    "sample": matrix[:5],
                }
            )

        return raw_pois, invalid_items

    def _table_to_matrix(self, table: Any) -> List[List[str]]:
        matrix: List[List[str]] = []

        for row in table.rows:
            cells = [self._clean_cell_text(cell.text) for cell in row.cells]
            while cells and cells[-1] == "":
                cells.pop()
            if any(cells):
                matrix.append(cells)

        return matrix

    def _clean_cell_text(self, text: str) -> str:
        text = text or ""
        text = text.replace("\xa0", " ")
        text = text.replace("\u3000", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _parse_horizontal_table(
        self,
        matrix: List[List[str]],
        table_index: int,
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        raw_pois: List[Dict[str, Any]] = []
        invalid_items: List[Dict[str, Any]] = []

        if len(matrix) < 2:
            return raw_pois, invalid_items

        headers = matrix[0]
        header_map = self._build_header_map(headers)

        if "name" not in header_map:
            invalid_items.append({"table_index": table_index, "reason": "横向表格缺少名称列，尝试其他解析方式"})
            return [], invalid_items

        for row_index, cells in enumerate(matrix[1:], start=2):
            if not any(cells):
                continue

            item: Dict[str, Any] = {
                "_source": "docx_table_horizontal",
                "_source_table_index": table_index,
                "_source_row_index": row_index,
            }

            for field_name, header_index in header_map.items():
                if header_index < len(cells):
                    item[field_name] = cells[header_index]

            if item.get("name"):
                raw_pois.append(item)
            else:
                invalid_items.append(
                    {
                        "table_index": table_index,
                        "row_index": row_index,
                        "reason": "该行缺少景点名称，已跳过",
                        "raw": cells,
                    }
                )

        return raw_pois, invalid_items

    def _parse_vertical_table(
        self,
        matrix: List[List[str]],
        table_index: int,
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        raw_pois: List[Dict[str, Any]] = []
        invalid_items: List[Dict[str, Any]] = []

        field_rows: List[Tuple[str, List[str]]] = []

        for row_index, row in enumerate(matrix, start=1):
            if len(row) < 2:
                continue

            field_name = self._map_header_name(row[0])
            if not field_name:
                continue

            values = row[1:]
            field_rows.append((field_name, values))

        if not field_rows:
            return raw_pois, invalid_items

        fields = {field for field, _ in field_rows}
        if "name" not in fields:
            invalid_items.append({"table_index": table_index, "reason": "纵向表格缺少景点名称字段，尝试其他解析方式"})
            return [], invalid_items

        max_items = max(len(values) for _, values in field_rows)
        for item_index in range(max_items):
            item: Dict[str, Any] = {
                "_source": "docx_table_vertical",
                "_source_table_index": table_index,
                "_source_item_index": item_index + 1,
            }

            for field_name, values in field_rows:
                if item_index < len(values):
                    value = values[item_index]
                    if value:
                        item[field_name] = value

            if item.get("name"):
                raw_pois.append(item)

        return raw_pois, invalid_items

    def _parse_flat_block_table(
        self,
        matrix: List[List[str]],
        table_index: int,
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        兼容当前上传文件这种结构。

        你这个 docx 在解析后常见形态是：
        表头：
            景区名称
            景点ID
            景点名称
            具体位置
            建筑/景观参数
            核心功能
            文化内涵
            详细介绍
            游玩亮点
            演艺/开放信息
            备注

        然后数据不是一行一个 POI，而是按字段顺序依次展开：
            灵山胜境
            LS-001
            灵山大照壁
            ...
            备注内容

            灵山胜境
            LS-002
            五明桥
            ...
            备注内容
        """
        raw_pois: List[Dict[str, Any]] = []
        invalid_items: List[Dict[str, Any]] = []

        tokens = self._flatten_table_tokens(matrix)
        if not tokens:
            return raw_pois, invalid_items

        all_items: List[Dict[str, Any]] = []

        # 一个 docx 表里可能同时有“子表1”和“表2”，所以要按表头切分多段
        header_positions = self._find_header_positions(tokens)

        if not header_positions:
            return raw_pois, invalid_items

        for segment_index, header_start in enumerate(header_positions):
            next_header_start = header_positions[segment_index + 1] if segment_index + 1 < len(header_positions) else len(tokens)
            segment_tokens = tokens[header_start:next_header_start]

            segment_items, segment_invalid = self._parse_flat_segment(
                segment_tokens=segment_tokens,
                table_index=table_index,
                segment_index=segment_index,
            )
            all_items.extend(segment_items)
            invalid_items.extend(segment_invalid)

        raw_pois.extend(all_items)
        return raw_pois, invalid_items

    def _flatten_table_tokens(self, matrix: List[List[str]]) -> List[str]:
        tokens: List[str] = []

        for row in matrix:
            for cell in row:
                cell = self._clean_cell_text(cell)
                if not cell:
                    continue

                # 有些 Word 表格会把多行内容塞进同一个 cell
                parts = [part.strip() for part in re.split(r"\n+", cell) if part.strip()]
                tokens.extend(parts)

        # 去掉完全重复的连续 token，避免合并单元格导致重复
        compacted: List[str] = []
        for token in tokens:
            if compacted and compacted[-1] == token:
                continue
            compacted.append(token)

        return compacted

    def _find_header_positions(self, tokens: List[str]) -> List[int]:
        positions: List[int] = []

        for index, token in enumerate(tokens):
            mapped = self._map_header_name(token)
            if mapped != "scenic_name":
                continue

            window = tokens[index : index + 20]
            mapped_window = [self._map_header_name(item) for item in window]
            mapped_set = {item for item in mapped_window if item}

            if "poi_id" in mapped_set and "name" in mapped_set:
                positions.append(index)

        return positions

    def _parse_flat_segment(
        self,
        segment_tokens: List[str],
        table_index: int,
        segment_index: int,
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        raw_pois: List[Dict[str, Any]] = []
        invalid_items: List[Dict[str, Any]] = []

        header_fields: List[str] = []
        header_raw: List[str] = []
        header_end = 0

        for index, token in enumerate(segment_tokens):
            field = self._map_header_name(token)

            if field:
                header_fields.append(field)
                header_raw.append(token)
                header_end = index + 1
                continue

            # 表头读取中断：已经读到 name 之后，遇到非表头，认为数据开始
            if "name" in header_fields and len(header_fields) >= 3:
                break

            # 还没读够表头时，跳过类似“子表1：...”这种标题行
            if token.startswith("子表") or token.startswith("表"):
                header_end = index + 1
                continue

        if not header_fields or "name" not in header_fields:
            invalid_items.append(
                {
                    "table_index": table_index,
                    "segment_index": segment_index,
                    "reason": "未识别到扁平表格表头或表头缺少景点名称",
                    "sample": segment_tokens[:20],
                }
            )
            return raw_pois, invalid_items

        data_tokens = segment_tokens[header_end:]
        data_tokens = [token for token in data_tokens if not self._looks_like_section_title(token)]

        field_count = len(header_fields)
        if field_count <= 0:
            return raw_pois, invalid_items

        index = 0
        item_no = 1

        while index + field_count <= len(data_tokens):
            block = data_tokens[index : index + field_count]

            # 如果 block 中又出现了一整组表头，跳过
            mapped_block = [self._map_header_name(token) for token in block]
            if "poi_id" in mapped_block and "name" in mapped_block and "scenic_name" in mapped_block:
                index += field_count
                continue

            item: Dict[str, Any] = {
                "_source": "docx_table_flat_block",
                "_source_table_index": table_index,
                "_source_segment_index": segment_index,
                "_source_item_index": item_no,
            }

            for field_name, value in zip(header_fields, block):
                value = self._clean_cell_text(value)
                if value:
                    item[field_name] = value

            if self._is_probable_poi_item(item):
                item = self._postprocess_docx_item(item)
                raw_pois.append(item)
                item_no += 1
                index += field_count
            else:
                invalid_items.append(
                    {
                        "table_index": table_index,
                        "segment_index": segment_index,
                        "item_index": item_no,
                        "reason": "该数据块不像有效 POI，尝试向后滑动一格继续解析",
                        "raw": block,
                    }
                )
                index += 1

        return raw_pois, invalid_items

    def _looks_like_section_title(self, text: str) -> bool:
        text = (text or "").strip()
        if not text:
            return True
        return bool(
            re.match(r"^子表\s*\d+", text)
            or re.match(r"^表\s*\d+", text)
            or text.endswith("景点数据集")
            or text.endswith("数据集")
            or text in {"数据集说明", "字段规范"}
        )

    def _is_probable_poi_item(self, item: Dict[str, Any]) -> bool:
        name = str(item.get("name") or "").strip()
        poi_id = str(item.get("poi_id") or "").strip()
        scenic_name = str(item.get("scenic_name") or "").strip()

        if not name:
            return False

        # 典型 ID：LS-001、NH-001
        if poi_id and re.match(r"^[A-Za-z]{1,8}[-_]\d{1,5}$", poi_id):
            return True

        # 没有 ID 也允许，只要有景区名 + 名称
        if scenic_name and name:
            return True

        return False

    def _postprocess_docx_item(self, item: Dict[str, Any]) -> Dict[str, Any]:
        """
        把你文档里的字段合并成后端标准字段。
        """
        item = dict(item)

        scenic_name = str(item.get("scenic_name") or "").strip()
        name = str(item.get("name") or "").strip()
        address = str(item.get("address") or "").strip()

        params = str(item.get("params") or "").strip()
        function = str(item.get("function") or "").strip()
        culture = str(item.get("culture") or "").strip()
        intro = str(item.get("intro") or "").strip()
        tags = str(item.get("tags") or "").strip()
        opening_info = str(item.get("opening_info") or "").strip()
        remark = str(item.get("remark") or "").strip()

        intro_parts = []
        if intro:
            intro_parts.append(intro)
        if culture:
            intro_parts.append(f"文化内涵：{culture}")
        if function:
            intro_parts.append(f"核心功能：{function}")
        if params:
            intro_parts.append(f"建筑/景观参数：{params}")
        if opening_info:
            intro_parts.append(f"演艺/开放信息：{opening_info}")
        if remark:
            intro_parts.append(f"备注：{remark}")

        if intro_parts:
            item["intro"] = self._trim_text(" ".join(intro_parts), 300)
        elif name:
            item["intro"] = f"{name}是{scenic_name or '景区'}内的重要点位。"

        if tags:
            item["tags"] = tags
        elif function:
            item["tags"] = function

        if address:
            item["address"] = address
        elif scenic_name and name:
            item["address"] = f"{scenic_name}{name}"

        if not item.get("point_type"):
            item["point_type"] = normalize_point_type(function, name)

        if not item.get("stay_minutes"):
            item["stay_minutes"] = self._guess_stay_minutes(name=name, intro=item.get("intro", ""), point_type=item.get("point_type", ""))

        return item

    def _guess_stay_minutes(self, name: str, intro: str, point_type: str) -> int:
        text = f"{name} {intro} {point_type}"

        if any(keyword in text for keyword in ["演出", "表演", "动态", "博览馆", "梵宫", "禅寺", "坛城"]):
            return 30
        if any(keyword in text for keyword in ["广场", "大道", "桥", "照壁", "浮雕", "塔", "柱"]):
            return 15
        if any(keyword in text for keyword in ["卫生间", "厕所", "服务中心", "停车场"]):
            return 5
        return 10

    def parse_docx_paragraphs(self, document: Document) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        text = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
        return self.parse_txt_or_md(text=text, source="docx_paragraph")

    # ============================================================
    # JSON / CSV / TXT 原有解析
    # ============================================================

    def parse_json(self, text: str) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], str | None]:
        try:
            payload = json.loads(text)
        except json.JSONDecodeError as exc:
            raise ValueError(f"JSON 解析失败: {exc.msg}") from exc

        scenic_name = None
        if isinstance(payload, list):
            raw_items = payload
        elif isinstance(payload, dict):
            scenic_name = payload.get("scenic_name")
            raw_items = payload.get("pois")
            if raw_items is None and isinstance(payload.get("data"), list):
                raw_items = payload["data"]
            if raw_items is None:
                raw_items = [payload]
        else:
            raise ValueError("JSON 文件格式错误，需为数组或包含 pois 的对象")

        raw_pois: List[Dict[str, Any]] = []
        invalid_items: List[Dict[str, Any]] = []

        for item in raw_items or []:
            if not isinstance(item, dict):
                invalid_items.append({"raw": item, "reason": "JSON 项不是对象，已跳过"})
                continue
            copied = dict(item)
            copied.setdefault("_source", "json")
            raw_pois.append(copied)

        return raw_pois, invalid_items, scenic_name

    def parse_csv(self, text: str) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        try:
            reader = csv.DictReader(io.StringIO(text))
        except csv.Error as exc:
            raise ValueError(f"CSV 解析失败: {exc}") from exc

        raw_pois: List[Dict[str, Any]] = []
        invalid_items: List[Dict[str, Any]] = []

        for row_index, row in enumerate(reader, start=2):
            if not row or not any((value or "").strip() for value in row.values()):
                continue

            normalized_row: Dict[str, Any] = {"_source": "csv", "_source_row_index": row_index}

            for field_name, raw_value in row.items():
                mapped_field = self._map_header_name(field_name or "")
                normalized_row[mapped_field or field_name] = (raw_value or "").strip()

            if not normalized_row.get("name"):
                invalid_items.append({"row_index": row_index, "reason": "CSV 缺少名称列，已跳过"})
                continue

            raw_pois.append(normalized_row)

        return raw_pois, invalid_items

    def parse_txt_or_md(self, text: str, source: str = "text_rule") -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        raw_pois: List[Dict[str, Any]] = []
        invalid_items: List[Dict[str, Any]] = []

        chunks = [item.strip(" \r\n\t-") for item in re.split(r"\n{2,}|\r\n\r\n", text) if item.strip()]
        if len(chunks) <= 1:
            chunks = [line.strip(" \r\n\t-") for line in text.splitlines() if line.strip()]

        for chunk in chunks:
            name = self._extract_text_name(chunk)
            lat, lng = self._extract_coordinates(chunk)

            if not name:
                invalid_items.append({"raw_text": chunk, "reason": "缺少名称，已跳过"})
                continue

            raw_pois.append(
                {
                    "_source": source,
                    "name": name,
                    "aliases": self._extract_aliases(chunk),
                    "lat": lat,
                    "lng": lng,
                    "point_type": normalize_point_type("", name),
                    "stay_minutes": self._extract_stay_minutes(chunk) or 10,
                    "intro": self._extract_intro(chunk, name=name),
                    "address": self._extract_address(chunk),
                    "tags": self._extract_tags(chunk),
                }
            )

        return raw_pois, invalid_items

    # ============================================================
    # 标准化 POI
    # ============================================================

    def normalize_poi(
        self,
        raw: Dict[str, Any],
        scenic_id: str,
        scenic_name: str | None,
        index: int,
        use_tencent_geocode: bool = True,
        city: str | None = None,
        address_hint: str | None = None,
    ) -> Dict[str, Any]:
        source = str(raw.get("_source") or "unknown")
        name = str(raw.get("name") or raw.get("title") or "").strip()

        if not name:
            return {"valid": False, "invalid": {"raw": raw, "reason": "缺少 name"}}

        poi_id = str(raw.get("poi_id") or raw.get("id") or "").strip() or f"{scenic_id.upper()}-POI-{index:03d}"
        aliases = self._normalize_aliases(raw.get("aliases"))
        tags = self._normalize_tags(raw.get("tags"))
        point_type = normalize_point_type(str(raw.get("point_type") or raw.get("type") or ""), name)
        stay_minutes = self._to_int(raw.get("stay_minutes") or raw.get("duration")) or 10
        address = str(raw.get("address") or "").strip()
        intro = self._build_intro(raw=raw, name=name, address=address)

        lat = self._to_float(raw.get("lat") or raw.get("latitude"))
        lng = self._to_float(raw.get("lng") or raw.get("lon") or raw.get("longitude"))

        # 支持 "经纬度" / "坐标" 这种字段
        if lat is None or lng is None:
            extracted_lat, extracted_lng = self._extract_coordinates_from_raw(raw)
            lat = lat if lat is not None else extracted_lat
            lng = lng if lng is not None else extracted_lng

        geocode_status = "original" if lat is not None and lng is not None else "skipped"
        confidence = self._base_confidence(source)

        if lat is None or lng is None:
            if use_tencent_geocode:
                enriched = self.enrich_poi_with_tencent_location(
                    name=name,
                    scenic_name=scenic_name or str(raw.get("scenic_name") or ""),
                    city=city,
                    address=address,
                    address_hint=address_hint,
                )
                if enriched:
                    lat = enriched["lat"]
                    lng = enriched["lng"]
                    address = address or enriched.get("address", "")
                    geocode_status = "tencent_resolved"
                    confidence = max(confidence, 0.82)
                else:
                    geocode_status = "failed"
            else:
                geocode_status = "skipped"

        if lat is None or lng is None:
            return {
                "valid": False,
                "invalid": {
                    "poi_id": poi_id,
                    "name": name,
                    "source": source,
                    "address": address,
                    "reason": "缺少合法坐标，且腾讯补全失败" if geocode_status == "failed" else "缺少合法坐标",
                    "geocode_status": geocode_status,
                },
            }

        poi = {
            "poi_id": poi_id,
            "name": name,
            "aliases": aliases,
            "lat": lat,
            "lng": lng,
            "point_type": point_type or "poi",
            "stay_minutes": stay_minutes,
            "intro": intro,
            "address": address,
            "tags": tags,
            "source": source,
            "geocode_status": geocode_status,
            "confidence": round(confidence, 2),
        }

        return {"valid": True, "poi": poi}

    def _extract_coordinates_from_raw(self, raw: Dict[str, Any]) -> Tuple[float | None, float | None]:
        for key in ["坐标", "经纬度", "location", "coordinate", "coordinates"]:
            value = raw.get(key)
            if not value:
                continue
            lat, lng = self._extract_coordinates(str(value))
            if lat is not None and lng is not None:
                return lat, lng
        return None, None

    def enrich_poi_with_tencent_location(
        self,
        name: str,
        scenic_name: str | None = None,
        city: str | None = None,
        address: str | None = None,
        address_hint: str | None = None,
    ) -> Dict[str, Any] | None:
        if self.tencent_map_client is None:
            return None

        try:
            return self.tencent_map_client.resolve_poi_location(
                name=name,
                scenic_name=scenic_name,
                city=city,
                address_hint=address_hint or address,
            )
        except Exception:
            return None

    def save_pois(
        self,
        scenic_id: str,
        scenic_name: str | None,
        pois: List[Dict[str, Any]],
        overwrite: bool = True,
        invalid_items: List[Dict[str, Any]] | None = None,
        source_file: str | None = None,
    ) -> Dict[str, Any]:
        invalid_items = invalid_items or []

        meta = {
            "source_file": source_file,
            "imported_at": datetime.now(timezone.utc).isoformat(),
            "poi_count": len(pois),
            "valid_count": len(pois),
            "invalid_count": len(invalid_items),
        }

        saved = self.repository.save_pois(
            scenic_id=scenic_id,
            scenic_name=scenic_name,
            pois=pois,
            overwrite=overwrite,
            meta=meta,
        )

        return {
            "scenic_id": saved["scenic_id"],
            "scenic_name": saved.get("scenic_name"),
            "poi_count": len(saved["pois"]),
            "pois": saved["pois"],
            "invalid_count": len(invalid_items),
            "invalid_items": invalid_items,
            "meta": meta,
            "message": f"已导入 {len(saved['pois'])} 个 POI",
        }

    def load_pois(self, scenic_id: str) -> List[Dict[str, Any]]:
        return self.repository.get_pois(scenic_id)

    # ============================================================
    # 字段映射 / 工具函数
    # ============================================================

    def _build_header_map(self, headers: List[str]) -> Dict[str, int]:
        mapped: Dict[str, int] = {}

        for index, header in enumerate(headers):
            field_name = self._map_header_name(header)
            if field_name and field_name not in mapped:
                mapped[field_name] = index

        return mapped

    def _map_header_name(self, header: str) -> str | None:
        raw = (header or "").strip()
        normalized = re.sub(r"[\s_\-：:（）()/\\]+", "", raw).lower()

        if not normalized:
            return None

        for target_field, aliases in FIELD_ALIASES.items():
            for alias in aliases:
                alias_normalized = re.sub(r"[\s_\-：:（）()/\\]+", "", alias).lower()
                if normalized == alias_normalized:
                    return target_field

        for target_field, aliases in FIELD_ALIASES.items():
            for alias in aliases:
                alias_normalized = re.sub(r"[\s_\-：:（）()/\\]+", "", alias).lower()
                if alias_normalized and (alias_normalized in normalized or normalized in alias_normalized):
                    return target_field

        return None

    @staticmethod
    def _normalize_aliases(value: Any) -> List[str]:
        if value is None:
            return []
        if isinstance(value, list):
            return [str(item).strip() for item in value if str(item).strip()]
        if isinstance(value, str):
            return [item.strip() for item in re.split(r"[|,，/、；;\s]+", value) if item.strip()]
        return []

    @staticmethod
    def _normalize_tags(value: Any) -> List[str]:
        if value is None:
            return []
        if isinstance(value, list):
            return [str(item).strip() for item in value if str(item).strip()]
        if isinstance(value, str):
            return [item.strip() for item in re.split(r"[|,，/、；;\s]+", value) if item.strip()]
        return []

    def _build_intro(self, raw: Dict[str, Any], name: str, address: str) -> str:
        intro = str(raw.get("intro") or raw.get("description") or raw.get("desc") or "").strip()

        if intro:
            return self._trim_text(intro, 300)

        if address:
            return self._trim_text(f"{name}位于{address}", 300)

        return self._trim_text(f"{name}是景区内的重点点位。", 300)

    @staticmethod
    def _to_float(value: Any) -> float | None:
        if value is None or value == "":
            return None

        text = str(value).strip()
        text = text.replace("纬度", "").replace("经度", "")
        text = text.replace("：", ":").strip()

        match = re.search(r"-?\d+(?:\.\d+)?", text)
        if not match:
            return None

        try:
            return float(match.group(0))
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _to_int(value: Any) -> int | None:
        if value is None or value == "":
            return None

        match = re.search(r"\d+", str(value))
        if not match:
            return None

        try:
            return int(match.group(0))
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _trim_text(text: str, max_length: int) -> str:
        text = re.sub(r"\s+", " ", (text or "")).strip()
        if len(text) <= max_length:
            return text
        return text[: max_length - 1].rstrip() + "…"

    @staticmethod
    def _base_confidence(source: str) -> float:
        if source in {"docx_table_horizontal", "docx_table_vertical", "docx_table_flat_block", "json", "csv"}:
            return 0.95
        if source in {"docx_paragraph", "text_rule"}:
            return 0.65
        return 0.6

    @staticmethod
    def _extract_text_name(text: str) -> str | None:
        first_line = text.splitlines()[0].strip()
        first_line = re.sub(r"^(?:#+|\*|-|\d+\.)\s*", "", first_line)

        for separator in [":", "：", "（", "(", "。", "，", ","]:
            if separator in first_line:
                candidate = first_line.split(separator, 1)[0].strip()
                if candidate:
                    return candidate[:40]

        fallback = re.match(r"^([^\s]{1,40})", first_line)
        return fallback.group(1).strip() if fallback else None

    @staticmethod
    def _extract_coordinates(text: str) -> Tuple[float | None, float | None]:
        patterns = [
            r"(?:坐标|经纬度|位置)[:：]?\s*([0-9]+(?:\.[0-9]+)?)\s*[,，]\s*([0-9]+(?:\.[0-9]+)?)",
            r"\b([0-9]+(?:\.[0-9]+)?)\s*[,，]\s*([0-9]+(?:\.[0-9]+)?)\b",
            r"lat[:：]?\s*([0-9]+(?:\.[0-9]+)?).*?(?:lng|lon|longitude)[:：]?\s*([0-9]+(?:\.[0-9]+)?)",
        ]

        for pattern in patterns:
            match = re.search(pattern, text, flags=re.I | re.S)
            if match:
                first = float(match.group(1))
                second = float(match.group(2))

                # 中国经纬度一般：纬度 3x，经度 1xx
                if first > 90 and second <= 90:
                    return second, first

                return first, second

        return None, None

    @staticmethod
    def _extract_stay_minutes(text: str) -> int | None:
        match = re.search(r"(?:建议停留|停留|游玩约|推荐游览)\s*([0-9]{1,3})\s*分钟", text)
        return int(match.group(1)) if match else None

    @staticmethod
    def _extract_aliases(text: str) -> List[str]:
        match = re.search(r"(?:别名|aliases?)[:：]\s*([^\n。；;]+)", text, flags=re.I)
        if not match:
            return []
        return [item.strip() for item in re.split(r"[|,，/、；;]", match.group(1)) if item.strip()]

    @staticmethod
    def _extract_address(text: str) -> str:
        match = re.search(r"(?:地址|位置|具体位置)[:：]\s*([^\n。]+)", text)
        return match.group(1).strip() if match else ""

    @staticmethod
    def _extract_tags(text: str) -> List[str]:
        match = re.search(r"(?:标签|特色|适合人群)[:：]\s*([^\n。]+)", text)
        if not match:
            return []
        return [item.strip() for item in re.split(r"[|,，/、；;]", match.group(1)) if item.strip()]

    def _extract_intro(self, text: str, name: str) -> str:
        intro = re.sub(r"^(?:#+\s*)?" + re.escape(name) + r"\s*[:：]?", "", text).strip()
        intro = re.sub(r"(?:坐标|经纬度|位置)[:：]?\s*[0-9]+\.[0-9]+\s*[,，]\s*[0-9]+\.[0-9]+", "", intro)
        intro = re.sub(r"(?:建议停留|停留|游玩约|推荐游览)\s*[0-9]{1,3}\s*分钟", "", intro)
        intro = re.sub(r"(?:别名|aliases?)[:：]\s*[^\n。；;]+", "", intro, flags=re.I)
        intro = re.sub(r"\s+", " ", intro).strip(" ，。；;")
        return self._trim_text(intro or f"{name}是景区内的重要点位。", 300)

    @staticmethod
    def _guess_scenic_name_from_docx(document: Document) -> str | None:
        for paragraph in document.paragraphs[:10]:
            text = paragraph.text.strip()
            if not text:
                continue
            if "灵山胜境" in text:
                return "灵山胜境"
            if len(text) <= 40 and "数据集" not in text:
                return text
        return None
