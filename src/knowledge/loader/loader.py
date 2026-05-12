from pathlib import Path
from typing import Dict, Optional

from docx import Document as DocxDocument

from ..chunker.schema import Document


class TextLoader:
    SUPPORTED_SUFFIXES = {".txt", ".md", ".docx"}

    def load(self, file_path: str, extra_metadata: Optional[Dict] = None) -> Document:
        path = Path(file_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_SUFFIXES:
            raise ValueError(f"暂不支持的文件类型: {suffix}")

        if suffix == ".docx":
            text = self._load_docx(path)
        else:
            text = path.read_text(encoding="utf-8")

        metadata = {
            "file_name": path.name,
            "file_path": str(path),
            "file_type": suffix.lstrip("."),
            "source": path.stem,
        }

        if extra_metadata:
            metadata.update(extra_metadata)

        return Document(text=text, metadata=metadata)

    @staticmethod
    def _load_docx(path: Path) -> str:
        doc = DocxDocument(str(path))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n\n".join(parts).strip()
