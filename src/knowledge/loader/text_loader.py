from docx import Document as DocxDocument
from knowledge.chunker.schema import Document


class TextLoader:

    def load(self, file_path: str) -> Document:
        doc = DocxDocument(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        return Document(
            text=text,
            metadata={
                "source": "docx",
                "file_path": file_path
            }
        )